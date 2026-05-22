import os
import datetime
from sqlalchemy.orm import Session
from app.models.book import Book
from app.models.chapter import Chapter


def get_output_path(book: Book, chapter: Chapter) -> str:
    base = book.output_folder or os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "output")
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    book_dir = os.path.join(base, book.title)
    date_dir = os.path.join(book_dir, date_str)
    os.makedirs(date_dir, exist_ok=True)

    ext = book.output_format or "md"
    filename = f"第{chapter.chapter_number}章-{chapter.title}.{ext}"
    if chapter.title:
        filename = f"{chapter.chapter_number:03d}-{chapter.title}.{ext}"
    else:
        filename = f"chapter_{chapter.chapter_number:03d}.{ext}"
    return os.path.join(date_dir, filename)


def export_chapter(book: Book, chapter: Chapter, db: Session) -> str:
    file_path = get_output_path(book, chapter)
    fmt = book.output_format or "md"

    if fmt == "docx":
        _export_docx(chapter.content, file_path)
    else:
        with open(file_path, "w", encoding="utf-8") as f:
            if fmt == "md":
                f.write(f"# {chapter.title}\n\n")
            f.write(chapter.content)
    return file_path


def _export_docx(content: str, file_path: str):
    from docx import Document
    doc = Document()
    for line in content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(file_path)


def export_all_chapters(book: Book, db: Session):
    chapters = db.query(Chapter).filter(Chapter.book_id == book.id).order_by(Chapter.chapter_number).all()
    base = book.output_folder or os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "output")
    book_dir = os.path.join(base, book.title)
    os.makedirs(book_dir, exist_ok=True)

    fmt = book.output_format or "md"
    full_path = os.path.join(book_dir, f"全书-{book.title}.{fmt}")

    with open(full_path, "w", encoding="utf-8") as f:
        for ch in chapters:
            if fmt == "md":
                f.write(f"# {ch.title}\n\n")
            f.write(ch.content)
            f.write("\n\n---\n\n")
    return full_path